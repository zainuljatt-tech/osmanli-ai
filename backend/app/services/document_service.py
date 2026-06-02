from sqlalchemy.orm import Session
from sqlalchemy import select, desc, text
from app.models.document import Document, DocumentChunk, DocumentStatus, DocumentType
from app.services.ai_service import ai_service
from app.core.config import settings
from typing import Optional
import os
import uuid
import json
import logging
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads"


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
        os.makedirs(UPLOAD_DIR, exist_ok=True)

    def upload_document(self, user_id: str, filename: str, content: bytes) -> Document:
        file_ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        ext_map = {"pdf": DocumentType.PDF, "docx": DocumentType.DOCX, "txt": DocumentType.TXT, "csv": DocumentType.CSV}
        doc_type = ext_map.get(file_ext, DocumentType.TXT)

        file_id = str(uuid.uuid4())
        safe_filename = f"{file_id}_{filename}"
        file_path = os.path.join(UPLOAD_DIR, safe_filename)

        with open(file_path, "wb") as f:
            f.write(content)

        doc = Document(
            user_id=user_id, filename=safe_filename, original_filename=filename,
            file_type=doc_type, file_size=len(content), storage_path=file_path,
            status=DocumentStatus.PROCESSING,
        )
        self.db.add(doc)
        self.db.flush()

        try:
            chunks = self._process_document(doc, file_path)
            for i, chunk_text in enumerate(chunks):
                embedding = ai_service.generate_embedding(chunk_text)
                embedding_value = json.dumps(embedding) if settings.USE_SQLITE and embedding else embedding
                self.db.add(DocumentChunk(
                    document_id=doc.id, chunk_index=i, content=chunk_text, embedding=embedding_value,
                ))
            doc.status = DocumentStatus.READY
            doc.chunk_count = len(chunks)
            self.db.flush()
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            doc.status = DocumentStatus.FAILED
            self.db.flush()

        return doc

    def _process_document(self, doc: Document, file_path: str) -> list[str]:
        try:
            if doc.file_type == DocumentType.PDF:
                reader = PdfReader(file_path)
                return self._chunk_text("\n".join(p.extract_text() for p in reader.pages))
            elif doc.file_type == DocumentType.DOCX:
                docx = DocxDocument(file_path)
                return self._chunk_text("\n".join(p.text for p in docx.paragraphs))
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    return self._chunk_text(f.read())
        except Exception as e:
            logger.error(f"Process error: {e}")
            return ["[Error processing document]"]

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        if not text.strip():
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            if end < len(text):
                end = text.rfind(" ", start, end)
                if end == -1 or end <= start:
                    end = min(start + chunk_size, len(text))
            chunks.append(text[start:end].strip())
            start = end - overlap
            if start < 0:
                start = 0
        return [c for c in chunks if c]

    def search_documents(self, user_id: str, query: str, document_ids: Optional[list[str]] = None, limit: int = 5) -> list[dict]:
        if not settings.USE_SQLITE:
            query_embedding = ai_service.generate_embedding(query)
            if query_embedding and not all(v == 0.0 for v in query_embedding):
                embedding_str = "[" + ",".join(str(v) for v in query_embedding) + "]"
                doc_filter = ""
                params = {"user_id": user_id, "query_emb": embedding_str, "query_emb2": embedding_str, "limit": limit}
                if document_ids:
                    placeholders = ",".join([f"'{d}'" for d in document_ids])
                    doc_filter = f" AND dc.document_id IN ({placeholders})"
                rows = self.db.execute(text(f"""
                    SELECT dc.id, dc.document_id, dc.content, dc.chunk_index,
                           1 - (dc.embedding <=> :query_emb) as score, d.original_filename
                    FROM document_chunks dc JOIN documents d ON d.id = dc.document_id
                    WHERE d.user_id = :user_id AND dc.embedding IS NOT NULL {doc_filter}
                    ORDER BY dc.embedding <=> :query_emb2 LIMIT :limit
                """), params).fetchall()
                return [{"chunk_id": str(r[0]), "document_id": str(r[1]), "content": r[2],
                         "chunk_index": r[3], "score": float(r[4]) if r[4] else 0, "filename": r[5]} for r in rows]

        q = select(DocumentChunk).join(Document).where(
            Document.user_id == user_id, DocumentChunk.content.ilike(f"%{query}%")
        )
        if document_ids:
            q = q.where(DocumentChunk.document_id.in_(document_ids))
        chunks = self.db.execute(q.limit(limit)).scalars().all()
        return [{"chunk_id": str(c.id), "document_id": str(c.document_id), "content": c.content,
                 "chunk_index": c.chunk_index, "score": 1.0, "filename": ""} for c in chunks]

    def get_user_documents(self, user_id: str) -> list[Document]:
        return list(self.db.execute(
            select(Document).where(Document.user_id == user_id).order_by(desc(Document.created_at))
        ).scalars().all())

    def delete_document(self, doc_id: str, user_id: str):
        doc = self.db.execute(select(Document).where(Document.id == doc_id, Document.user_id == user_id)).scalar_one_or_none()
        if doc:
            if os.path.exists(doc.storage_path):
                os.remove(doc.storage_path)
            self.db.delete(doc)
            self.db.flush()


def get_document_service(db: Session) -> DocumentService:
    return DocumentService(db)
